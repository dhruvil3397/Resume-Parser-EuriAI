from docx import Document
from io import BytesIO


def txt_to_docx_bytes(text: str) -> bytes:
    """
    Convert plain text into a .docx file represented as bytes.

    Rules:
    - Preserves line breaks: each input line becomes a paragraph.
    - Blank lines are kept as empty paragraphs.
    - Avoids trailing whitespace issues by using line.strip() to detect empties only.

    Args:
        text: Input plain text.

    Returns:
        Bytes of a .docx file suitable for saving or returning from an API.
    """
    doc = Document()
    for line in text.splitlines():
        if line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()
